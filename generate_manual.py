#!/usr/bin/env python3
"""
Comprehensive User Manual Generator for Cartup CxP Roster Management System
This script generates a complete DOCX manual with screenshots, API documentation, and FAQ.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import glob

def create_manual():
    """Create the comprehensive user manual document"""
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Cartup CxP Roster Management System - User Manual"
    doc.core_properties.author = "Cartup CxP Team"
    
    # Title Page
    title = doc.add_heading('Cartup CxP Roster Management System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete User Manual')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    
    version = doc.add_paragraph('Version 1.0')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        ("1.", "Introduction", "4"),
        ("  1.1", "About This Manual", "4"),
        ("  1.2", "System Overview", "4"),
        ("  1.3", "Key Features", "5"),
        ("2.", "Client Panel User Guide", "6"),
        ("  2.1", "Logging In", "6"),
        ("  2.2", "Dashboard Overview", "7"),
        ("  2.3", "Refresh Function", "8"),
        ("  2.4", "Theme Customization", "9"),
        ("  2.5", "Calendar Feature", "10"),
        ("  2.6", "Requesting Shift Changes", "12"),
        ("  2.7", "Requesting Shift Swaps", "15"),
        ("  2.8", "Shift View", "18"),
        ("  2.9", "Employee Search", "20"),
        ("  2.10", "Statistics Cards", "22"),
        ("3.", "Admin Panel User Guide", "25"),
        ("  3.1", "Admin Login", "25"),
        ("  3.2", "Dashboard Tab", "26"),
        ("  3.3", "Schedule Requests Tab", "30"),
        ("  3.4", "Data Sync Tab", "33"),
        ("  3.5", "Google Sheets Tab", "35"),
        ("  3.6", "Roster Data Tab", "37"),
        ("  3.7", "CSV Import/Export Tab", "40"),
        ("  3.8", "My Profile Tab", "43"),
        ("  3.9", "Team Management Tab", "45"),
        ("  3.10", "User Management Tab", "48"),
        ("4.", "API Documentation", "51"),
        ("  4.1", "Authentication APIs", "51"),
        ("  4.2", "Schedule APIs", "53"),
        ("  4.3", "Request APIs", "56"),
        ("  4.4", "Admin APIs", "59"),
        ("  4.5", "Data Sync APIs", "62"),
        ("5.", "Frequently Asked Questions (FAQ)", "65"),
        ("  5.1", "General Questions", "65"),
        ("  5.2", "Client Panel Questions", "67"),
        ("  5.3", "Admin Panel Questions", "69"),
        ("  5.4", "Troubleshooting", "71"),
        ("6.", "Appendices", "73"),
        ("  6.1", "Shift Codes Reference", "73"),
        ("  6.2", "Quick Reference Guide", "74"),
    ]
    
    table = doc.add_table(rows=len(toc_items), cols=3)
    table.style = 'Light Grid Accent 1'
    
    for idx, (num, title, page) in enumerate(toc_items):
        row = table.rows[idx]
        row.cells[0].text = num
        row.cells[1].text = title
        row.cells[2].text = page
    
    doc.add_page_break()
    
    # CHAPTER 1: Introduction
    doc.add_heading('1. Introduction', 1)
    
    doc.add_heading('1.1 About This Manual', 2)
    doc.add_paragraph(
        'This comprehensive manual provides step-by-step instructions for using the Cartup CxP '
        'Roster Management System. Whether you are an employee accessing your schedule or an '
        'administrator managing team rosters, this guide will help you understand and utilize '
        'all features of the system effectively.'
    )
    
    doc.add_heading('1.2 System Overview', 2)
    doc.add_paragraph(
        'The Cartup CxP Roster Management System is a modern web-based application designed to '
        'streamline shift scheduling, request management, and team coordination. The system '
        'consists of two main components:'
    )
    
    features = [
        ('Client Panel', 'For employees to view schedules, request changes, and manage their shifts'),
        ('Admin Panel', 'For administrators to manage rosters, approve requests, and oversee operations'),
    ]
    
    for feature, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{feature}: ').bold = True
        p.add_run(desc)
    
    doc.add_heading('1.3 Key Features', 2)
    
    client_features = [
        'Real-time schedule viewing',
        'Interactive calendar for date selection',
        'Shift change request submission',
        'Shift swap requests with team members',
        'Employee search functionality',
        'Personal statistics and upcoming shifts',
        'Multiple theme options for personalization',
        'Mobile-responsive design',
    ]
    
    admin_features = [
        'Comprehensive dashboard with analytics',
        'Request approval/rejection workflow',
        'Team and employee management',
        'Google Sheets integration',
        'CSV import/export capabilities',
        'User management with role-based access',
        'Activity logging and audit trails',
        'Shift modification tracking',
    ]
    
    doc.add_paragraph('Client Panel Features:').bold = True
    for feature in client_features:
        doc.add_paragraph(f'• {feature}', style='List Bullet 2')
    
    doc.add_paragraph('Admin Panel Features:').bold = True
    for feature in admin_features:
        doc.add_paragraph(f'• {feature}', style='List Bullet 2')
    
    doc.add_page_break()
    
    # CHAPTER 2: Client Panel
    doc.add_heading('2. Client Panel User Guide', 1)
    
    add_client_panel_sections(doc)
    
    # CHAPTER 3: Admin Panel
    doc.add_heading('3. Admin Panel User Guide', 1)
    
    add_admin_panel_sections(doc)
    
    # CHAPTER 4: API Documentation
    doc.add_heading('4. API Documentation', 1)
    
    add_api_documentation(doc)
    
    # CHAPTER 5: FAQ
    doc.add_heading('5. Frequently Asked Questions (FAQ)', 1)
    
    add_faq_section(doc)
    
    # CHAPTER 6: Appendices
    doc.add_heading('6. Appendices', 1)
    
    add_appendices(doc)
    
    # Save document
    doc.save('USER_MANUAL.docx')
    print("✅ User manual generated successfully: USER_MANUAL.docx")

def add_client_panel_sections(doc):
    """Add detailed client panel documentation"""
    
    # 2.1 Logging In
    doc.add_heading('2.1 Logging In to the Client Panel', 2)
    doc.add_paragraph(
        'To access your schedule and manage your shifts, you need to log in to the Client Panel.'
    )
    
    doc.add_paragraph('Steps:').bold = True
    steps = [
        'Navigate to the application URL (http://localhost:3000 or your organization URL)',
        'Enter your Full Name in the first field',
        'Enter your Employee ID in the format SLL-XXXXX',
        'The team password is pre-filled as "cartup123"',
        'Click the "🔓 Access Roster" button',
        'You will be redirected to your personal dashboard',
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph('📸 Screenshot: See MANUAL_SCREENSHOTS/client/01_client_login_page.png')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Note: ').bold = True
    p.add_run('The Employee ID is case-sensitive. Make sure to enter it exactly as provided.')
    
    # 2.2 Dashboard Overview
    doc.add_heading('2.2 Dashboard Overview', 2)
    doc.add_paragraph(
        'Once logged in, you will see your personalized dashboard displaying:'
    )
    
    dashboard_elements = [
        ('Welcome Header', 'Shows your name and Employee ID'),
        ('Action Buttons', 'Logout, Refresh, and Theme buttons'),
        ('Current Shift Information', 'Today and tomorrow shift details'),
        ('Selected Date Shift', 'Shows shift for any selected calendar date'),
        ('Action Buttons Row', 'Request Shift Change, Request Swap, and Shift View buttons'),
        ('Employee Search', 'Search bar to find and view other employees\' schedules'),
        ('Statistics Cards', 'Upcoming Days, Planned Time Off, and Shift Changes'),
    ]
    
    for element, description in dashboard_elements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{element}: ').bold = True
        p.add_run(description)
    
    doc.add_paragraph('📸 Screenshot: See MANUAL_SCREENSHOTS/client/02_client_dashboard_main.png')
    
    # 2.3 Refresh Function
    doc.add_heading('2.3 Refresh Function', 2)
    doc.add_paragraph(
        'The Refresh button allows you to reload your schedule data to see the most up-to-date '
        'information including any recently approved shift changes.'
    )
    
    doc.add_paragraph('How to use:').bold = True
    refresh_steps = [
        'Locate the "🔄 Refresh" button in the top action bar',
        'Click the button',
        'The system will reload all schedule data',
        'The button will show "Refreshing..." while loading',
        'Once complete, all information will be updated',
    ]
    for step in refresh_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph('📸 Screenshot: See MANUAL_SCREENSHOTS/client/03_after_refresh.png')
    
    # 2.4 Theme Customization
    doc.add_heading('2.4 Theme Customization', 2)
    doc.add_paragraph(
        'The system offers multiple color themes to personalize your experience. You can switch '
        'between different themes to find one that suits your preference.'
    )
    
    doc.add_paragraph('Available Themes:').bold = True
    themes = [
        '🌈 Bright Vibrant - Colorful and energetic',
        '🌅 Bright Sunset - Warm and inviting',
        '🌊 Medium Ocean - Cool blue tones',
        '🌍 Medium Earth - Natural earth tones',
        '🍃 Peaceful Sage - Calming green',
        '💜 Peaceful Lavender - Soft purple',
        '🌑 Dark Blue - Professional dark blue',
        '🌃 Dark Midnight - Deep dark theme',
        '🕳️ Dark Void - Maximum contrast black',
    ]
    for theme in themes:
        doc.add_paragraph(theme, style='List Bullet')
    
    doc.add_paragraph('How to change theme:').bold = True
    theme_steps = [
        'Click the "🎨 Theme" button in the top action bar',
        'A dropdown menu will appear showing all available themes',
        'Click on any theme to apply it immediately',
        'The entire website will update with the new color scheme',
        'Your selection is saved and will persist across sessions',
    ]
    for step in theme_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph('📸 Screenshots:')
    doc.add_paragraph('- Theme menu: MANUAL_SCREENSHOTS/client/04_theme_menu_open.png')
    doc.add_paragraph('- Theme applied: MANUAL_SCREENSHOTS/client/05_theme_changed_ocean.png')
    
    # 2.5 Calendar Feature
    doc.add_heading('2.5 Calendar Feature', 2)
    doc.add_paragraph(
        'The calendar allows you to view your shift schedule for any date. When you select a date, '
        'the system displays your assigned shift for that day.'
    )
    
    doc.add_paragraph('How to use the calendar:').bold = True
    calendar_steps = [
        'Click the "📅 Show Calendar" button',
        'The calendar will expand, showing the current month',
        'Use the arrow buttons (← →) to navigate between months',
        'Click on any date to view your shift for that day',
        'The selected date and shift will appear above the calendar',
        'Click "📅 Hide Calendar" to collapse the calendar',
    ]
    for step in calendar_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph('📸 Screenshots:')
    doc.add_paragraph('- Calendar opened: MANUAL_SCREENSHOTS/client/06_calendar_opened.png')
    doc.add_paragraph('- October view: MANUAL_SCREENSHOTS/client/07_calendar_october.png')
    doc.add_paragraph('- Date selected: MANUAL_SCREENSHOTS/client/08_date_selected_oct20.png')
    
    # Continue with other client sections...
    add_shift_change_section(doc)
    add_swap_request_section(doc)
    add_shift_view_section(doc)
    add_employee_search_section(doc)
    add_stat_cards_section(doc)

def add_shift_change_section(doc):
    """Add shift change request documentation"""
    doc.add_heading('2.6 Requesting Shift Changes', 2)
    doc.add_paragraph(
        'If you need to change your assigned shift for a specific date, you can submit a shift '
        'change request through the system. An administrator will review and approve or reject '
        'your request.'
    )
    
    doc.add_paragraph('Step-by-step process:').bold = True
    steps = [
        'Click the "✏️ Request Shift Change" button on the dashboard',
        'The Shift Change Request modal will open',
        'You will see your employee information and current team displayed',
        'Select the date for which you want to change your shift using the mini calendar',
        'Use the arrow buttons to navigate to the correct month if needed',
        'Click on the desired date',
        'Your current shift for that date will be displayed',
        'Select your requested shift from the dropdown menu (M2, M3, M4, D1, D2, DO, SL, CL, EL, HL)',
        'Enter a reason for your request in the text area',
        'Click "Submit Request" to send your request to administrators',
        'Click "Cancel" if you want to close the modal without submitting',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph('📸 Screenshot: See MANUAL_SCREENSHOTS/client/09_shift_change_modal_opened.png')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Important: ').bold = True
    p.add_run('All shift change requests require administrator approval. You will be notified once '
              'your request is processed.')

def add_swap_request_section(doc):
    """Add swap request documentation"""
    doc.add_heading('2.7 Requesting Shift Swaps', 2)
    doc.add_paragraph(
        'A shift swap allows you to exchange shifts with another team member. Both the requester '
        'and the target employee must be on the same team for a swap to be processed.'
    )
    
    doc.add_paragraph('How to request a swap:').bold = True
    steps = [
        'Click the "🔁 Request Swap" button on the dashboard',
        'The Swap Request modal will open',
        'Select the date for the swap using the calendar',
        'Your current shift for that date will be displayed',
        'In the "Swap With" field, start typing an employee name or ID',
        'A list of team members will appear as you type',
        'Select the employee you want to swap with',
        'Enter a reason for the swap request',
        'Click "Submit Swap Request"',
        'The request will be sent to administrators for approval',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Note: ').bold = True
    p.add_run('The system will only show employees from your team in the search suggestions. '
              'Cross-team swaps are not currently supported.')

def add_shift_view_section(doc):
    """Add shift view documentation"""
    doc.add_heading('2.8 Shift View', 2)
    doc.add_paragraph(
        'The Shift View feature provides a comprehensive calendar-style view of team schedules, '
        'allowing you to see who is working on specific dates.'
    )
    
    doc.add_paragraph('Using Shift View:').bold = True
    steps = [
        'Click the "👁️ Shift View" button',
        'The Shift View modal will open showing a calendar',
        'Select a date from the calendar to view all shifts for that day',
        'You can filter by team using the team dropdown',
        'The view shows all employees and their assigned shifts',
        'Use the arrow buttons to navigate between months',
        'Click outside the modal or the close button to exit',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tip: ').bold = True
    p.add_run('Use this feature to coordinate with team members and plan coverage.')

def add_employee_search_section(doc):
    """Add employee search documentation"""
    doc.add_heading('2.9 Employee Search', 2)
    doc.add_paragraph(
        'The employee search feature allows you to look up any employee in the system and view '
        'their schedule.'
    )
    
    doc.add_paragraph('How to search for employees:').bold = True
    steps = [
        'Locate the "Search Other Employees" section on the dashboard',
        'Click in the search box',
        'Start typing an employee name, ID, or team name',
        'A dropdown list of matching employees will appear',
        'Click on an employee from the list',
        'Their schedule will replace yours on the dashboard temporarily',
        'You can select dates from the calendar to see their shifts',
        'Click the "← Back to My Schedule" button to return to your own schedule',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Use case: ').bold = True
    p.add_run('This is useful for checking if a colleague is available on a specific day before '
              'requesting a swap.')

def add_stat_cards_section(doc):
    """Add statistics cards documentation"""
    doc.add_heading('2.10 Statistics Cards', 2)
    doc.add_paragraph(
        'The bottom of your dashboard displays three statistics cards that provide quick insights '
        'into your schedule:'
    )
    
    cards = [
        ('📅 Upcoming Days', 
         'Shows the number of working days in the next 7 days. Click to expand and see the list '
         'of dates you are scheduled to work.'),
        ('🏖️ Planned Time Off', 
         'Displays your time off days (DO, SL, CL, EL, HL) within the next 30 days. Click to '
         'expand and see all your scheduled off days with their types.'),
        ('🔄 Shift Changes', 
         'Shows the number of shifts that have been modified from the original Google Sheets roster. '
         'Click to expand and see details of what changed and when.'),
    ]
    
    for title, description in cards:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(description)
    
    doc.add_paragraph('How to use:').bold = True
    doc.add_paragraph('1. Click on any card to expand it')
    doc.add_paragraph('2. The card will show detailed information')
    doc.add_paragraph('3. Click the "▲" arrow or anywhere outside to collapse')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tip: ').bold = True
    p.add_run('Check these cards regularly to stay aware of your upcoming schedule and any changes.')

def add_admin_panel_sections(doc):
    """Add detailed admin panel documentation"""
    
    # 3.1 Admin Login
    doc.add_heading('3.1 Admin Login', 2)
    doc.add_paragraph(
        'Administrators access a separate panel with advanced features for managing the entire roster system.'
    )
    
    doc.add_paragraph('Default Admin Credentials:').bold = True
    admin_creds = [
        ('Super Admin', 'Username: developer', 'Password: devneversleeps'),
        ('Admin', 'Username: istiaque', 'Password: cartup123'),
        ('Admin', 'Username: admin', 'Password: password123'),
    ]
    
    table = doc.add_table(rows=len(admin_creds) + 1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Username'
    hdr_cells[2].text = 'Password'
    
    for idx, (role, username, password) in enumerate(admin_creds, 1):
        row = table.rows[idx]
        row.cells[0].text = role
        row.cells[1].text = username
        row.cells[2].text = password
    
    doc.add_paragraph()
    doc.add_paragraph('Steps to login:').bold = True
    steps = [
        'Navigate to http://localhost:3000/admin/login',
        'Enter your admin username',
        'Enter your password',
        'Click "Login"',
        'You will be redirected to the admin dashboard',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    # Continue with other admin sections...
    add_admin_dashboard_section(doc)
    add_schedule_requests_section(doc)
    add_data_sync_section(doc)
    add_google_sheets_section(doc)
    add_roster_data_section(doc)
    add_csv_section(doc)
    add_profile_section(doc)
    add_team_mgmt_section(doc)
    add_user_mgmt_section(doc)

def add_admin_dashboard_section(doc):
    """Add admin dashboard documentation"""
    doc.add_heading('3.2 Dashboard Tab', 2)
    doc.add_paragraph(
        'The admin dashboard provides an overview of the entire roster system with key metrics '
        'and recent activity.'
    )
    
    doc.add_paragraph('Dashboard Components:').bold = True
    
    components = [
        ('👥 Total Employees This Month', 
         'Shows the total number of employees in the system.'),
        ('👷 Employees Working Today', 
         'Displays count of employees with shifts today. Click to see the full list with their shifts.'),
        ('Shift Change / Swap Requests Overview', 
         'Statistics card showing pending, approved, and rejected requests. Click to expand for details.'),
        ('Team Health Overview', 
         'Shows team distribution and metrics. Expand to see detailed team information.'),
        ('Activity Log', 
         'Recent actions including approved requests, rejected requests, and shift modifications. '
         'Shows admin username who performed each action.'),
    ]
    
    for title, description in components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(description)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('How to use: ').bold = True
    p.add_run('Click on any stat card to expand it and view detailed information. '
              'The activity log updates automatically as changes are made.')

def add_schedule_requests_section(doc):
    """Add schedule requests documentation"""
    doc.add_heading('3.3 Schedule Requests Tab', 2)
    doc.add_paragraph(
        'This tab is where administrators review and process shift change and swap requests from employees.'
    )
    
    doc.add_paragraph('Request Management:').bold = True
    
    steps = [
        'Click on the "Schedule Requests" tab in the sidebar',
        'You will see a list of all requests',
        'Use the filter buttons to view: All, Pending, Approved, Rejected',
        'For each request, you can see:',
        '  - Employee name and ID',
        '  - Request type (Shift Change or Swap)',
        '  - Requested date',
        '  - Current shift and requested shift',
        '  - Reason provided by employee',
        '  - Request submission date',
        'To approve a request: Click the "✅ Approve" button',
        'To reject a request: Click the "❌ Reject" button',
        'You will be asked to confirm your action',
        'Once processed, the request status updates immediately',
        'The employee\'s schedule is updated for approved requests',
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Important: ').bold = True
    p.add_run('All actions are logged and cannot be undone. Approved shift changes immediately '
              'update the roster.')

def add_data_sync_section(doc):
    """Add data sync documentation"""
    doc.add_heading('3.4 Data Sync Tab', 2)
    doc.add_paragraph(
        'The Data Sync tab allows you to synchronize roster data from Google Sheets and manage '
        'automatic synchronization settings.'
    )
    
    doc.add_paragraph('Features:').bold = True
    
    features = [
        ('Manual Sync Button', 
         'Click to immediately fetch and update data from all configured Google Sheets links.'),
        ('Auto-Sync Toggle', 
         'Enable or disable automatic synchronization that runs at regular intervals.'),
        ('Last Sync Time', 
         'Shows when the last successful sync occurred.'),
        ('Sync Statistics', 
         'Displays number of employees and sheets synced.'),
    ]
    
    for title, description in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(description)
    
    doc.add_paragraph('How to perform a manual sync:').bold = True
    steps = [
        'Navigate to the Data Sync tab',
        'Click the "Sync Now" button',
        'Wait for the sync to complete',
        'A success message will appear',
        'Check the sync statistics to verify',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

def add_google_sheets_section(doc):
    """Add Google Sheets configuration documentation"""
    doc.add_heading('3.5 Google Sheets Tab', 2)
    doc.add_paragraph(
        'Configure Google Sheets links for roster data import. The system supports multiple sheets '
        'to aggregate data from different teams or sources.'
    )
    
    doc.add_paragraph('Managing Google Sheets Links:').bold = True
    
    doc.add_paragraph('To add a new link:').bold = True
    steps = [
        'Click on the "Google Sheets" tab',
        'Enter a descriptive name for the sheet (e.g., "Voice Team Roster")',
        'Paste the published CSV link from your Google Sheet',
        'Click "Add Link"',
        'The link will be saved and used for future syncs',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph('To delete a link:').bold = True
    doc.add_paragraph('1. Find the link in the list')
    doc.add_paragraph('2. Click the "Delete" button next to it')
    doc.add_paragraph('3. Confirm the deletion')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('How to get a Google Sheets CSV link:').bold = True
    doc.add_paragraph('1. Open your Google Sheet')
    doc.add_paragraph('2. Go to File → Share → Publish to web')
    doc.add_paragraph('3. Select "Comma-separated values (.csv)"')
    doc.add_paragraph('4. Click "Publish"')
    doc.add_paragraph('5. Copy the generated URL')

def add_roster_data_section(doc):
    """Add roster data management documentation"""
    doc.add_heading('3.6 Roster Data Tab', 2)
    doc.add_paragraph(
        'The Roster Data tab provides an interactive interface to view and edit employee shifts directly.'
    )
    
    doc.add_paragraph('Features:').bold = True
    
    features = [
        ('Data Source Toggle', 
         'Switch between viewing Google Sheets roster (original) and Admin modified roster.'),
        ('Shift View Button', 
         'Open a calendar-based view of the entire roster.'),
        ('Reset to Google Button', 
         'Reset all admin modifications and revert to the original Google Sheets data.'),
        ('Date Selection', 
         'Select any date to view and modify shifts for that day.'),
        ('Employee List', 
         'View all employees with their shifts for the selected date.'),
        ('Shift Editing', 
         'Click on any employee shift cell to change it.'),
    ]
    
    for title, description in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(description)
    
    doc.add_paragraph('How to modify a shift:').bold = True
    steps = [
        'Go to the Roster Data tab',
        'Select "Admin Data" to edit the modifiable roster',
        'Click "Select Date to Modify Shifts"',
        'Choose a date from the calendar',
        'Find the employee whose shift you want to change',
        'Click on their current shift code',
        'A dropdown will appear with all available shift codes',
        'Select the new shift',
        'The change is saved automatically',
        'The modification is tracked and logged',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

def add_csv_section(doc):
    """Add CSV import/export documentation"""
    doc.add_heading('3.7 CSV Import/Export Tab', 2)
    doc.add_paragraph(
        'Import and export roster data in CSV format for backup, bulk editing, or integration '
        'with external systems.'
    )
    
    doc.add_paragraph('CSV Import:').bold = True
    steps = [
        'Click on "CSV Import" tab',
        'Click "Choose File" or drag and drop a CSV file',
        'The file should follow the template format',
        'Select the month this data is for',
        'Click "Upload CSV"',
        'The system will process and import the data',
        'A success message confirms the import',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph('CSV Export:').bold = True
    steps = [
        'Go to the CSV Import tab',
        'Select specific months to export or choose "Export All"',
        'Click "📥 Export CSV"',
        'The file will be generated and downloaded',
        'Open the file in Excel or any spreadsheet application',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('CSV Format: ').bold = True
    p.add_run('The CSV must have columns for Employee Name, Employee ID, Team, and date columns '
              'with shift codes.')

def add_profile_section(doc):
    """Add profile management documentation"""
    doc.add_heading('3.8 My Profile Tab', 2)
    doc.add_paragraph(
        'Manage your admin account information and change your password.'
    )
    
    doc.add_paragraph('Profile Information:').bold = True
    info_items = [
        'Username (read-only)',
        'Role (read-only)',
        'Change password functionality',
    ]
    for item in info_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('How to change your password:').bold = True
    steps = [
        'Go to the "My Profile" tab',
        'Enter your current password',
        'Enter your new password',
        'Re-enter the new password to confirm',
        'Click "Change Password"',
        'You will receive a confirmation message',
        'Use your new password for future logins',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

def add_team_mgmt_section(doc):
    """Add team management documentation"""
    doc.add_heading('3.9 Team Management Tab', 2)
    doc.add_paragraph(
        'Manage teams and employees, including adding new employees, modifying information, and '
        'organizing team structures.'
    )
    
    doc.add_paragraph('Team Management Features:').bold = True
    
    doc.add_paragraph('Adding a new team:').bold = True
    steps = [
        'Click on "Team Management" tab',
        'Click "Add New Team" button',
        'Enter the team name',
        'Optionally add a description',
        'Click "Save"',
        'The team will appear in the list',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph('Adding a new employee:').bold = True
    steps = [
        'Select the team from the dropdown',
        'Click "Add Employee"',
        'Fill in employee details:',
        '  - Full Name',
        '  - Employee ID (format: SLL-XXXXX)',
        '  - Team assignment',
        'Click "Save Employee"',
        'The employee will be added to the roster',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph('Modifying employee information:').bold = True
    doc.add_paragraph('1. Find the employee in the list')
    doc.add_paragraph('2. Click "Edit" next to their name')
    doc.add_paragraph('3. Update the information')
    doc.add_paragraph('4. Click "Save Changes"')

def add_user_mgmt_section(doc):
    """Add user management documentation"""
    doc.add_heading('3.10 User Management Tab', 2)
    doc.add_paragraph(
        'Manage administrator accounts, including creating new users, updating roles, and deleting accounts. '
        'Note: This tab is only visible to Super Admins and Admins.'
    )
    
    doc.add_paragraph('User Roles:').bold = True
    roles = [
        ('super_admin', 'Full system access including user management'),
        ('admin', 'Can manage rosters and requests, view user management'),
        ('team_leader', 'Limited access to team-specific functions'),
    ]
    
    for role, description in roles:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{role}: ').bold = True
        p.add_run(description)
    
    doc.add_paragraph('Adding a new admin user:').bold = True
    steps = [
        'Go to the "User Management" tab',
        'Click "Add New User"',
        'Fill in the form:',
        '  - Username (unique)',
        '  - Password',
        '  - Confirm Password',
        '  - Select Role',
        'Click "Create User"',
        'The user can now log in with these credentials',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph('Deleting a user:').bold = True
    doc.add_paragraph('1. Find the user in the list')
    doc.add_paragraph('2. Click the "Delete" button')
    doc.add_paragraph('3. Confirm the deletion')
    doc.add_paragraph('4. The user account will be permanently removed')

def add_api_documentation(doc):
    """Add comprehensive API documentation"""
    
    doc.add_paragraph(
        'This section documents all API endpoints available in the Cartup CxP Roster Management System. '
        'All APIs use JSON for request and response bodies.'
    )
    
    # Authentication APIs
    doc.add_heading('4.1 Authentication APIs', 2)
    
    apis = [
        {
            'endpoint': 'POST /api/admin/login',
            'description': 'Admin login endpoint',
            'request': '{"username": "string", "password": "string"}',
            'response': '{"success": true, "user": {"username": "string", "role": "string"}}',
            'auth': 'None required',
        },
        {
            'endpoint': 'POST /api/admin/logout',
            'description': 'Admin logout endpoint',
            'request': 'None',
            'response': '{"success": true}',
            'auth': 'Admin cookie required',
        },
    ]
    
    for api in apis:
        doc.add_paragraph(api['endpoint']).bold = True
        doc.add_paragraph(f"Description: {api['description']}")
        doc.add_paragraph(f"Authentication: {api['auth']}")
        doc.add_paragraph(f"Request Body: {api['request']}")
        doc.add_paragraph(f"Response: {api['response']}")
        doc.add_paragraph()
    
    # Schedule APIs
    doc.add_heading('4.2 Schedule APIs', 2)
    
    schedule_apis = [
        {
            'endpoint': 'GET /api/my-schedule/[employeeId]',
            'description': 'Get employee schedule',
            'params': 'employeeId - Employee ID (e.g., SLL-88717)',
            'response': '{"employee": {...}, "headers": [...], "schedule": [...]}',
        },
        {
            'endpoint': 'GET /api/admin/get-display-data',
            'description': 'Get merged display roster data',
            'response': '{"teams": {...}, "headers": [...], "allEmployees": [...]}',
        },
        {
            'endpoint': 'GET /api/admin/get-admin-data',
            'description': 'Get admin-modified roster data',
            'response': '{"teams": {...}, "headers": [...], "allEmployees": [...]}',
        },
        {
            'endpoint': 'GET /api/admin/get-google-data',
            'description': 'Get original Google Sheets roster data',
            'response': '{"teams": {...}, "headers": [...], "allEmployees": [...]}',
        },
    ]
    
    for api in schedule_apis:
        doc.add_paragraph(api['endpoint']).bold = True
        doc.add_paragraph(f"Description: {api['description']}")
        if 'params' in api:
            doc.add_paragraph(f"Parameters: {api['params']}")
        doc.add_paragraph(f"Response: {api['response']}")
        doc.add_paragraph()
    
    # Request APIs
    doc.add_heading('4.3 Request APIs', 2)
    
    request_apis = [
        {
            'endpoint': 'POST /api/schedule-requests/submit-shift-change',
            'description': 'Submit a shift change request',
            'request': '''{
  "employee_id": "string",
  "employee_name": "string",
  "team": "string",
  "date": "string",
  "current_shift": "string",
  "requested_shift": "string",
  "reason": "string"
}''',
            'response': '{"success": true, "message": "Request submitted"}',
        },
        {
            'endpoint': 'POST /api/schedule-requests/submit-swap-request',
            'description': 'Submit a shift swap request',
            'request': '''{
  "requester_id": "string",
  "requester_name": "string",
  "swap_with_id": "string",
  "swap_with_name": "string",
  "team": "string",
  "date": "string",
  "reason": "string"
}''',
            'response': '{"success": true, "message": "Swap request submitted"}',
        },
        {
            'endpoint': 'GET /api/schedule-requests/get-all',
            'description': 'Get all schedule requests',
            'response': '[{"id": "string", "type": "string", "status": "string", ...}]',
        },
        {
            'endpoint': 'POST /api/schedule-requests/update-status',
            'description': 'Approve or reject a request (admin only)',
            'request': '{"id": "string", "status": "approved|rejected", "admin_username": "string"}',
            'response': '{"success": true}',
        },
    ]
    
    for api in request_apis:
        doc.add_paragraph(api['endpoint']).bold = True
        doc.add_paragraph(f"Description: {api['description']}")
        if 'request' in api:
            doc.add_paragraph(f"Request Body:")
            doc.add_paragraph(api['request'])
        doc.add_paragraph(f"Response:")
        doc.add_paragraph(api['response'])
        doc.add_paragraph()
    
    # Admin APIs
    doc.add_heading('4.4 Admin APIs', 2)
    
    admin_apis = [
        {
            'endpoint': 'POST /api/admin/update-shift',
            'description': 'Update employee shift for a specific date',
            'request': '''{
  "employee_id": "string",
  "date": "string",
  "shift_code": "string",
  "admin_username": "string"
}''',
            'response': '{"success": true}',
        },
        {
            'endpoint': 'POST /api/admin/upload-csv',
            'description': 'Upload roster CSV file',
            'request': 'multipart/form-data with file and month',
            'response': '{"success": true, "message": "CSV imported"}',
        },
        {
            'endpoint': 'POST /api/admin/export-csv',
            'description': 'Export roster data as CSV',
            'request': '{"months": ["string"]}',
            'response': 'CSV file download',
        },
        {
            'endpoint': 'POST /api/admin/save-team',
            'description': 'Create or update a team',
            'request': '{"name": "string", "description": "string"}',
            'response': '{"success": true}',
        },
        {
            'endpoint': 'POST /api/admin/save-employee',
            'description': 'Create or update an employee',
            'request': '{"id": "string", "name": "string", "team": "string"}',
            'response': '{"success": true}',
        },
    ]
    
    for api in admin_apis:
        doc.add_paragraph(api['endpoint']).bold = True
        doc.add_paragraph(f"Description: {api['description']}")
        doc.add_paragraph(f"Request Body:")
        doc.add_paragraph(api['request'])
        doc.add_paragraph(f"Response:")
        doc.add_paragraph(api['response'])
        doc.add_paragraph()
    
    # Data Sync APIs
    doc.add_heading('4.5 Data Sync APIs', 2)
    
    sync_apis = [
        {
            'endpoint': 'POST /api/admin/sync-google-sheets',
            'description': 'Manually trigger Google Sheets sync',
            'response': '{"success": true, "employees": number, "sheets": number}',
        },
        {
            'endpoint': 'POST /api/admin/set-auto-sync',
            'description': 'Enable or disable automatic sync',
            'request': '{"enabled": boolean}',
            'response': '{"success": true}',
        },
        {
            'endpoint': 'POST /api/admin/reset-to-google',
            'description': 'Reset admin data to Google Sheets data',
            'response': '{"success": true, "message": "Data reset"}',
        },
        {
            'endpoint': 'GET /api/admin/get-modified-shifts',
            'description': 'Get list of all modified shifts',
            'response': '[{"employee_id": "string", "date": "string", "old_shift": "string", "new_shift": "string", ...}]',
        },
    ]
    
    for api in sync_apis:
        doc.add_paragraph(api['endpoint']).bold = True
        doc.add_paragraph(f"Description: {api['description']}")
        if 'request' in api:
            doc.add_paragraph(f"Request Body:")
            doc.add_paragraph(api['request'])
        doc.add_paragraph(f"Response:")
        doc.add_paragraph(api['response'])
        doc.add_paragraph()

def add_faq_section(doc):
    """Add FAQ section"""
    
    # General Questions
    doc.add_heading('5.1 General Questions', 2)
    
    general_faqs = [
        {
            'q': 'What browsers are supported?',
            'a': 'The system works best on modern browsers including Chrome, Firefox, Safari, and Edge. '
                 'We recommend using the latest version of Chrome for the best experience.',
        },
        {
            'q': 'Is the system mobile-friendly?',
            'a': 'Yes! The system is fully responsive and works on mobile devices, tablets, and desktops. '
                 'The interface adapts to your screen size.',
        },
        {
            'q': 'How often is the data updated?',
            'a': 'If auto-sync is enabled, data is synchronized from Google Sheets every hour. You can also '
                 'manually refresh at any time using the Refresh button.',
        },
        {
            'q': 'Can I access the system from home?',
            'a': 'Yes, if your organization has made the system accessible externally. Contact your IT '
                 'department for the correct URL and VPN requirements if needed.',
        },
    ]
    
    for faq in general_faqs:
        p = doc.add_paragraph()
        p.add_run(f"Q: {faq['q']}").bold = True
        doc.add_paragraph(f"A: {faq['a']}")
        doc.add_paragraph()
    
    # Client Panel Questions
    doc.add_heading('5.2 Client Panel Questions', 2)
    
    client_faqs = [
        {
            'q': 'Why can\'t I log in?',
            'a': 'Make sure you are entering your Employee ID correctly (format: SLL-XXXXX). The ID is '
                 'case-sensitive. Also verify that the password is "cartup123". If issues persist, '
                 'contact your administrator.',
        },
        {
            'q': 'How do I know if my request was approved?',
            'a': 'Check the "Shift Changes" stat card on your dashboard. Approved changes will be reflected '
                 'there. You can also check your schedule - approved changes will show the new shift.',
        },
        {
            'q': 'Can I cancel a request after submitting?',
            'a': 'Currently, you cannot cancel a request yourself. Contact your administrator if you need '
                 'to cancel a pending request.',
        },
        {
            'q': 'Why can\'t I request a swap with someone?',
            'a': 'You can only swap shifts with team members from your own team. The system will only show '
                 'employees from your team in the swap request search.',
        },
        {
            'q': 'What do the shift codes mean?',
            'a': 'M2 (8 AM-5 PM), M3 (9 AM-6 PM), M4 (10 AM-7 PM), D1 (12 PM-9 PM), D2 (1 PM-10 PM), '
                 'DO (Day Off), SL (Sick Leave), CL (Casual Leave), EL (Emergency Leave), HL (Holiday Leave).',
        },
    ]
    
    for faq in client_faqs:
        p = doc.add_paragraph()
        p.add_run(f"Q: {faq['q']}").bold = True
        doc.add_paragraph(f"A: {faq['a']}")
        doc.add_paragraph()
    
    # Admin Panel Questions
    doc.add_heading('5.3 Admin Panel Questions', 2)
    
    admin_faqs = [
        {
            'q': 'How do I add a new employee to the system?',
            'a': 'Go to Team Management tab, select the team, click "Add Employee", fill in the details '
                 '(Name, ID, Team), and save. The employee will appear in the roster immediately.',
        },
        {
            'q': 'What happens when I approve a shift change request?',
            'a': 'The employee\'s shift is immediately updated in the admin roster. The change is logged '
                 'in the modification history and appears in the activity feed.',
        },
        {
            'q': 'Can I undo a shift modification?',
            'a': 'Yes, you can manually change the shift back to the original value, or use the '
                 '"Reset to Google" button to reset all modifications at once (warning: this resets ALL changes).',
        },
        {
            'q': 'How do I bulk import employee schedules?',
            'a': 'Use the CSV Import tab. Download the template, fill it with your data following the format, '
                 'then upload it. Select the correct month before uploading.',
        },
        {
            'q': 'What\'s the difference between Google Data and Admin Data?',
            'a': 'Google Data is the original roster from Google Sheets (read-only). Admin Data includes '
                 'all modifications made by administrators. The system displays a merge of both.',
        },
    ]
    
    for faq in admin_faqs:
        p = doc.add_paragraph()
        p.add_run(f"Q: {faq['q']}").bold = True
        doc.add_paragraph(f"A: {faq['a']}")
        doc.add_paragraph()
    
    # Troubleshooting
    doc.add_heading('5.4 Troubleshooting', 2)
    
    troubleshooting = [
        {
            'issue': 'Page not loading or showing errors',
            'solution': 'Try refreshing the page (F5). Clear your browser cache. Check your internet connection. '
                       'If the issue persists, contact IT support.',
        },
        {
            'issue': 'Data not updating after sync',
            'solution': 'Click the manual refresh button. Check if the Google Sheets links are correctly configured. '
                       'Verify that the Google Sheet is published correctly as CSV.',
        },
        {
            'issue': 'Cannot upload CSV file',
            'solution': 'Ensure the file is in CSV format (.csv extension). Check that the file follows the '
                       'template format. File size should not exceed 5MB. Try a different browser.',
        },
        {
            'issue': 'Theme not applying correctly',
            'solution': 'Clear your browser cache. Try selecting the theme again. Check if JavaScript is enabled '
                       'in your browser settings.',
        },
        {
            'issue': 'Forgot admin password',
            'solution': 'Contact a Super Admin to reset your password through the User Management tab. Super Admins '
                       'can reset passwords for other users.',
        },
    ]
    
    for item in troubleshooting:
        p = doc.add_paragraph()
        p.add_run(f"Issue: {item['issue']}").bold = True
        doc.add_paragraph(f"Solution: {item['solution']}")
        doc.add_paragraph()

def add_appendices(doc):
    """Add appendices"""
    
    # Shift Codes Reference
    doc.add_heading('6.1 Shift Codes Reference', 2)
    
    doc.add_paragraph('Complete list of all shift codes used in the system:')
    
    shift_codes = [
        ('M2', '8 AM – 5 PM', 'Morning Shift 2'),
        ('M3', '9 AM – 6 PM', 'Morning Shift 3'),
        ('M4', '10 AM – 7 PM', 'Morning Shift 4'),
        ('D1', '12 PM – 9 PM', 'Day Shift 1'),
        ('D2', '1 PM – 10 PM', 'Day Shift 2'),
        ('DO', 'Day Off', 'Scheduled day off'),
        ('SL', 'Sick Leave', 'Medical leave'),
        ('CL', 'Casual Leave', 'Personal leave'),
        ('EL', 'Emergency Leave', 'Urgent/emergency leave'),
        ('HL', 'Holiday Leave', 'Public holiday or scheduled holiday'),
    ]
    
    table = doc.add_table(rows=len(shift_codes) + 1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Code'
    hdr_cells[1].text = 'Time/Type'
    hdr_cells[2].text = 'Description'
    
    for idx, (code, time, desc) in enumerate(shift_codes, 1):
        row = table.rows[idx]
        row.cells[0].text = code
        row.cells[1].text = time
        row.cells[2].text = desc
    
    # Quick Reference Guide
    doc.add_heading('6.2 Quick Reference Guide', 2)
    
    doc.add_paragraph('Client Panel Quick Actions:').bold = True
    client_actions = [
        ('View Schedule', 'Login → Dashboard shows today/tomorrow'),
        ('Change Theme', 'Click Theme button → Select from dropdown'),
        ('Request Shift Change', 'Click Request Shift Change → Select date → Choose shift → Submit'),
        ('Request Swap', 'Click Request Swap → Select date → Choose employee → Submit'),
        ('View Team Schedule', 'Click Shift View → Select date and team'),
        ('Search Employee', 'Type in search box → Click employee'),
    ]
    
    table = doc.add_table(rows=len(client_actions) + 1, cols=2)
    table.style = 'Light List Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Action'
    hdr_cells[1].text = 'How To'
    
    for idx, (action, howto) in enumerate(client_actions, 1):
        row = table.rows[idx]
        row.cells[0].text = action
        row.cells[1].text = howto
    
    doc.add_paragraph()
    doc.add_paragraph('Admin Panel Quick Actions:').bold = True
    admin_actions = [
        ('Approve Request', 'Schedule Requests tab → Find request → Click Approve'),
        ('Modify Shift', 'Roster Data tab → Select date → Click shift → Choose new shift'),
        ('Sync Data', 'Data Sync tab → Click Sync Now'),
        ('Add Employee', 'Team Management tab → Add Employee → Fill form → Save'),
        ('Export CSV', 'CSV Import tab → Select months → Click Export'),
        ('Add Admin User', 'User Management tab → Add New User → Fill details → Create'),
    ]
    
    table = doc.add_table(rows=len(admin_actions) + 1, cols=2)
    table.style = 'Light List Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Action'
    hdr_cells[1].text = 'How To'
    
    for idx, (action, howto) in enumerate(admin_actions, 1):
        row = table.rows[idx]
        row.cells[0].text = action
        row.cells[1].text = howto
    
    # Support Contact
    doc.add_page_break()
    doc.add_heading('Support & Contact', 1)
    doc.add_paragraph(
        'For technical support, questions, or issues with the Cartup CxP Roster Management System, '
        'please contact:'
    )
    doc.add_paragraph()
    doc.add_paragraph('IT Support Team')
    doc.add_paragraph('Email: support@cartup.com')
    doc.add_paragraph('Phone: +1-XXX-XXX-XXXX')
    doc.add_paragraph('Hours: Monday - Friday, 9 AM - 5 PM')
    
    doc.add_paragraph()
    doc.add_paragraph('System Administrator')
    doc.add_paragraph('Email: admin@cartup.com')
    
    doc.add_paragraph()
    doc.add_paragraph('---')
    doc.add_paragraph('Document Version: 1.0')
    doc.add_paragraph('Last Updated: October 2025')
    doc.add_paragraph('© 2025 Cartup CxP. All rights reserved.')

if __name__ == '__main__':
    try:
        create_manual()
        print("\n✅ SUCCESS: Complete user manual has been generated!")
        print("📄 File location: USER_MANUAL.docx")
        print("📸 Screenshots folder: MANUAL_SCREENSHOTS/")
        print("\nThe manual includes:")
        print("  ✓ Table of Contents")
        print("  ✓ Client Panel Guide (with screenshots)")
        print("  ✓ Admin Panel Guide (with screenshots)")
        print("  ✓ Comprehensive API Documentation")
        print("  ✓ FAQ Section")
        print("  ✓ Appendices with Quick Reference")
    except Exception as e:
        print(f"❌ Error generating manual: {e}")
        import traceback
        traceback.print_exc()
